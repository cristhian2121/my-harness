from __future__ import annotations

import csv
import io
import math
import re
from hashlib import sha256
from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader
from qdrant_client import QdrantClient, models

from app.application.exceptions import UnsupportedDocumentFormatError
from app.domain.entities import (
    Document,
    DocumentChunk,
    ParsedDocument,
    RetrievedDocumentChunk,
)


class LocalDocumentStorage:
    def __init__(self, root_directory: Path) -> None:
        self._root_directory = root_directory
        self._root_directory.mkdir(parents=True, exist_ok=True)

    def store(self, *, document_id: str, filename: str, content: bytes) -> str:
        extension = Path(filename).suffix.lower()
        target_path = self._root_directory / f"{document_id}{extension}"
        target_path.write_bytes(content)
        return str(target_path)


class LocalDocumentParser:
    def __init__(self, *, text_chunk_size: int = 1200, row_chunk_size: int = 20) -> None:
        self._text_chunk_size = text_chunk_size
        self._row_chunk_size = row_chunk_size

    def parse(self, *, filename: str, content: bytes) -> ParsedDocument:
        extension = Path(filename).suffix.lower()
        if extension == ".pdf":
            return self._parse_pdf(content)
        if extension == ".csv":
            return self._parse_csv(content)
        if extension == ".docx":
            return self._parse_docx(content)
        if extension == ".xlsx":
            return self._parse_xlsx(content)
        if extension == ".doc":
            raise UnsupportedDocumentFormatError(
                "Legacy .doc files are not supported. Convert the file to .docx or .pdf."
            )
        if extension == ".xls":
            raise UnsupportedDocumentFormatError(
                "Legacy .xls files are not supported. Convert the file to .xlsx or .csv."
            )
        raise UnsupportedDocumentFormatError(
            "Unsupported document format. Supported formats: .pdf, .csv, .docx, .xlsx."
        )

    def _parse_pdf(self, content: bytes) -> ParsedDocument:
        reader = PdfReader(io.BytesIO(content))
        chunks: list[DocumentChunk] = []
        chunk_index = 0
        for page_number, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if not text:
                continue
            for chunk_text in self._split_text(text):
                chunks.append(
                    DocumentChunk(
                        id=None,
                        document_id="",
                        chunk_index=chunk_index,
                        text=chunk_text,
                        page_number=page_number,
                    )
                )
                chunk_index += 1
        return ParsedDocument(chunks=chunks, page_count=len(reader.pages))

    def _parse_csv(self, content: bytes) -> ParsedDocument:
        decoded = self._decode_text(content)
        reader = list(csv.reader(io.StringIO(decoded)))
        if not reader:
            return ParsedDocument(chunks=[])

        header = [cell.strip() for cell in reader[0]]
        rows = reader[1:]
        chunks: list[DocumentChunk] = []
        chunk_index = 0
        for start in range(0, len(rows), self._row_chunk_size):
            group = rows[start : start + self._row_chunk_size]
            if not group:
                continue
            row_start = start + 1
            row_end = start + len(group)
            lines = [self._format_tabular_row(header, row) for row in group]
            text = (
                f"Columns: {', '.join(header)}\n"
                f"Rows {row_start}-{row_end}:\n"
                + "\n".join(lines)
            ).strip()
            chunks.append(
                DocumentChunk(
                    id=None,
                    document_id="",
                    chunk_index=chunk_index,
                    text=text,
                    row_start=row_start,
                    row_end=row_end,
                    metadata={"columns": header},
                )
            )
            chunk_index += 1
        return ParsedDocument(chunks=chunks, row_count=len(rows))

    def _parse_docx(self, content: bytes) -> ParsedDocument:
        document = DocxDocument(io.BytesIO(content))
        sections: list[tuple[str | None, list[str]]] = []
        current_heading: str | None = None
        current_lines: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = getattr(paragraph.style, "name", "") or ""
            if style_name.lower().startswith("heading"):
                if current_lines:
                    sections.append((current_heading, current_lines))
                current_heading = text
                current_lines = []
                continue
            current_lines.append(text)

        if current_lines:
            sections.append((current_heading, current_lines))

        chunks: list[DocumentChunk] = []
        chunk_index = 0
        for heading, lines in sections or [(None, [p.text.strip() for p in document.paragraphs if p.text.strip()])]:
            section_text = "\n".join(lines).strip()
            if not section_text:
                continue
            for chunk_text in self._split_text(section_text):
                rendered_text = f"Heading: {heading}\n{chunk_text}" if heading else chunk_text
                chunks.append(
                    DocumentChunk(
                        id=None,
                        document_id="",
                        chunk_index=chunk_index,
                        text=rendered_text,
                        metadata={"heading": heading} if heading else None,
                    )
                )
                chunk_index += 1
        return ParsedDocument(chunks=chunks)

    def _parse_xlsx(self, content: bytes) -> ParsedDocument:
        workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        chunks: list[DocumentChunk] = []
        chunk_index = 0
        total_rows = 0
        sheet_names: list[str] = []

        for sheet_name in workbook.sheetnames:
            worksheet = workbook[sheet_name]
            rows = list(worksheet.iter_rows(values_only=True))
            if not rows:
                continue
            sheet_names.append(sheet_name)
            header = [self._normalize_cell(value) for value in rows[0]]
            data_rows = rows[1:]
            total_rows += len(data_rows)
            for start in range(0, len(data_rows), self._row_chunk_size):
                group = data_rows[start : start + self._row_chunk_size]
                if not group:
                    continue
                row_start = start + 1
                row_end = start + len(group)
                lines = [self._format_tabular_row(header, row) for row in group]
                text = (
                    f"Sheet: {sheet_name}\n"
                    f"Columns: {', '.join(header)}\n"
                    f"Rows {row_start}-{row_end}:\n"
                    + "\n".join(lines)
                ).strip()
                chunks.append(
                    DocumentChunk(
                        id=None,
                        document_id="",
                        chunk_index=chunk_index,
                        text=text,
                        sheet_name=sheet_name,
                        row_start=row_start,
                        row_end=row_end,
                        metadata={"columns": header},
                    )
                )
                chunk_index += 1
        return ParsedDocument(chunks=chunks, row_count=total_rows, sheet_names=sheet_names or None)

    def _split_text(self, text: str) -> list[str]:
        normalized = re.sub(r"\s+", " ", text).strip()
        if not normalized:
            return []
        if len(normalized) <= self._text_chunk_size:
            return [normalized]

        chunks: list[str] = []
        start = 0
        overlap = min(150, self._text_chunk_size // 5)
        while start < len(normalized):
            end = min(len(normalized), start + self._text_chunk_size)
            chunks.append(normalized[start:end].strip())
            if end == len(normalized):
                break
            start = max(end - overlap, start + 1)
        return chunks

    @staticmethod
    def _decode_text(content: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "latin-1"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        return content.decode("utf-8", errors="ignore")

    @staticmethod
    def _normalize_cell(value: object) -> str:
        if value is None:
            return ""
        return str(value).strip()

    def _format_tabular_row(self, header: list[str], row: tuple[object, ...] | list[str]) -> str:
        values = [self._normalize_cell(value) for value in row]
        if any(header):
            pairs = [f"{column or f'column_{index + 1}'}={values[index] if index < len(values) else ''}" for index, column in enumerate(header)]
            return "; ".join(pairs)
        return ", ".join(values)


class LocalHashEmbeddingService:
    def __init__(self, dimensions: int = 128) -> None:
        self._dimensions = dimensions

    def embed_texts(self, texts: list[str]) -> list[list[float]]:
        return [self._embed(text) for text in texts]

    def embed_query(self, text: str) -> list[float]:
        return self._embed(text)

    def _embed(self, text: str) -> list[float]:
        vector = [0.0] * self._dimensions
        tokens = re.findall(r"\w+", text.lower()) or [text.lower()]
        for token in tokens:
            digest = sha256(token.encode("utf-8")).digest()
            index = int.from_bytes(digest[:4], "big") % self._dimensions
            sign = -1.0 if digest[4] % 2 else 1.0
            weight = 1.0 + (digest[5] / 255.0)
            vector[index] += sign * weight

        norm = math.sqrt(sum(value * value for value in vector))
        if norm == 0:
            return vector
        return [value / norm for value in vector]


class LocalQdrantVectorStore:
    def __init__(self, *, root_directory: Path, collection_name: str = "document_chunks") -> None:
        self._root_directory = root_directory
        self._root_directory.mkdir(parents=True, exist_ok=True)
        self._collection_name = collection_name
        self._client: QdrantClient | None = None
        self._vector_size: int | None = None

    def index_chunks(
        self,
        *,
        username: str,
        document: Document,
        chunks: list[DocumentChunk],
        vectors: list[list[float]],
    ) -> None:
        if not chunks:
            return
        self._ensure_collection(len(vectors[0]))
        points = [
            models.PointStruct(
                id=int(chunk.id),
                vector=vectors[index],
                payload={
                    "username": username,
                    "document_id": document.id,
                    "filename": document.filename,
                    "text": chunk.text,
                    "chunk_index": chunk.chunk_index,
                    "page_number": chunk.page_number,
                    "sheet_name": chunk.sheet_name,
                    "row_start": chunk.row_start,
                    "row_end": chunk.row_end,
                },
            )
            for index, chunk in enumerate(chunks)
            if chunk.id is not None
        ]
        self._get_client().upsert(collection_name=self._collection_name, points=points, wait=True)

    def search(
        self,
        *,
        username: str,
        vector: list[float],
        document_ids: list[str],
        limit: int,
    ) -> list[RetrievedDocumentChunk]:
        if not self._ensure_collection_loaded(expected_vector_size=len(vector)):
            return []
        response = self._get_client().query_points(
            collection_name=self._collection_name,
            query=vector,
            query_filter=models.Filter(
                must=[
                    models.FieldCondition(
                        key="username",
                        match=models.MatchValue(value=username),
                    ),
                    models.FieldCondition(
                        key="document_id",
                        match=models.MatchAny(any=document_ids),
                    ),
                ]
            ),
            limit=limit,
            with_payload=True,
        )
        return [
            RetrievedDocumentChunk(
                chunk_id=int(hit.id),
                document_id=str(hit.payload["document_id"]),
                filename=str(hit.payload["filename"]),
                text=str(hit.payload["text"]),
                score=float(hit.score),
                chunk_index=int(hit.payload["chunk_index"]),
                page_number=self._optional_int(hit.payload.get("page_number")),
                sheet_name=self._optional_str(hit.payload.get("sheet_name")),
                row_start=self._optional_int(hit.payload.get("row_start")),
                row_end=self._optional_int(hit.payload.get("row_end")),
            )
            for hit in response.points
        ]

    def _ensure_collection(self, vector_size: int) -> None:
        if self._ensure_collection_loaded(expected_vector_size=vector_size):
            return

        self._get_client().create_collection(
            collection_name=self._collection_name,
            vectors_config=models.VectorParams(
                size=vector_size,
                distance=models.Distance.COSINE,
            ),
        )
        self._vector_size = vector_size

    def _get_client(self) -> QdrantClient:
        if self._client is None:
            self._client = QdrantClient(path=str(self._root_directory))
        return self._client

    def close(self) -> None:
        if self._client is None:
            return
        close = getattr(self._client, "close", None)
        if callable(close):
            close()
        self._client = None
        self._vector_size = None

    def _ensure_collection_loaded(self, *, expected_vector_size: int) -> bool:
        if self._vector_size is not None:
            if self._vector_size != expected_vector_size:
                raise ValueError(
                    "Existing vector collection size does not match the configured embedding size."
                )
            return True
        if not self._collection_exists():
            return False

        collection = self._get_client().get_collection(self._collection_name)
        loaded_size = self._extract_vector_size(collection)
        if loaded_size != expected_vector_size:
            raise ValueError(
                "Existing vector collection size does not match the configured embedding size."
            )
        self._vector_size = loaded_size
        return True

    def _collection_exists(self) -> bool:
        exists = getattr(self._get_client(), "collection_exists", None)
        if callable(exists):
            return bool(exists(collection_name=self._collection_name))
        try:
            self._get_client().get_collection(self._collection_name)
        except Exception:
            return False
        return True

    @staticmethod
    def _extract_vector_size(collection: object) -> int:
        vectors = collection.config.params.vectors
        if isinstance(vectors, dict):
            first_vector = next(iter(vectors.values()), None)
            size = getattr(first_vector, "size", None)
        else:
            size = getattr(vectors, "size", None)
        if size is None:
            raise ValueError("Unable to determine vector size for persisted collection.")
        return int(size)

    @staticmethod
    def _optional_int(value: object) -> int | None:
        if value is None:
            return None
        return int(value)

    @staticmethod
    def _optional_str(value: object) -> str | None:
        if value is None:
            return None
        text = str(value)
        return text or None
